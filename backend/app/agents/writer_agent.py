"""
Writer Agent - CV and Cover Letter Tailoring Specialist

This agent receives structured data from CV and Job agents and generates:
1. Tailored CV PDFs based on job requirements
2. Cover letter PDFs aligned with the job and company

The agent follows a human-in-the-loop workflow:
- Analyzes gaps between CV and job requirements
- Proposes modifications for user approval
- Generates content after approval
- Creates professional PDF outputs

Architecture:
- Input: Pre-processed JSON data from cv_agent (ResumeInfo) and job_agent (JobRequirements, CompanyInfo)
- Processing: Gap analysis → Tailoring plan → Content generation → PDF creation
- Output: Professional PDF documents (CV and cover letter)
- Pattern: LangChain agent with specialized tools for each stage

Workflow:
1. ANALYSIS: Compare CV against job requirements (analyze_cv_job_alignment)
2. REVIEW: Present tailoring plan and wait for user approval
3. GENERATION: Create tailored CV HTML (generate_tailored_cv_html)
4. APPROVAL: Show preview and wait for user confirmation
5. PDF CREATION: Generate final CV PDF (generate_cv_pdf)
6. COVER LETTER: Generate letter content and PDF (generate_cover_letter_content, generate_cover_letter_pdf)

Key Features:
- Human-in-the-loop at every critical step
- Preserves candidate's authentic voice
- Never fabricates content - only emphasizes/refines existing material
- Professional PDF generation with table-based CV layout
- Company-aware cover letter generation
"""

from langchain.agents import create_agent
from langchain_core.tools import tool
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from pathlib import Path
import json
import html
import dotenv
import os
import re
from typing import List, Dict, Any, Optional
from app.models.schemas import (
    CVTailoringPlan,
    CoverLetterContent,
    CompatibilityReport,
    JobSummary,
    SelectedContent,
    RewrittenContent,
)
from app.agents.scoring_agent import (
    calculate_semantic_similarity,
    calculate_bm25_score,
    calculate_compatibility_score,
    match_skill_pairs,
    assess_transferability_llm,
    calculate_compatibility_score_v2,
    build_gap_analysis,
    _calculate_compatibility_score_v2_internal,
)

dotenv.load_dotenv()

# ============================================
# WeasyPrint Import Helpers (Lazy Loading)
# ============================================

# Lazy import for WeasyPrint to avoid GTK dependency issues on Windows
# WeasyPrint will only be imported when PDF generation is actually needed
_weasyprint_available = None

def _check_weasyprint():
    """Check if WeasyPrint is available and can be imported."""
    global _weasyprint_available
    if _weasyprint_available is None:
        try:
            from weasyprint import HTML, CSS
            _weasyprint_available = True
        except (ImportError, OSError) as e:
            _weasyprint_available = False
            print(f"⚠️  WeasyPrint not available: {e}")
            print("PDF generation will not work. See installation instructions:")
            print("https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#windows")
    return _weasyprint_available

def _import_weasyprint():
    """Import WeasyPrint with error handling."""
    try:
        from weasyprint import HTML, CSS
        return HTML, CSS
    except (ImportError, OSError) as e:
        raise RuntimeError(
            f"WeasyPrint is not properly installed: {e}\n\n"
            "On Windows, WeasyPrint requires GTK libraries.\n"
            "Installation options:\n"
            "1. Install GTK via MSYS2: https://www.gtk.org/docs/installations/windows/\n"
            "2. Use WSL (Windows Subsystem for Linux)\n"
            "3. Use Docker\n"
            "4. Use an alternative PDF library (reportlab, fpdf)\n\n"
            "See: https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#windows"
        ) from e

# ============================================
# Constants
# ============================================

# Default output directory for generated PDFs
DEFAULT_OUTPUT_DIR = Path(__file__).parent.parent.parent / "data"

# CSS Template for CV (two-column table layout)
CV_CSS_TEMPLATE = """
@page {
    size: A4;
    margin: 2cm;
}

body {
    font-family: 'Helvetica', 'Arial', sans-serif;
    font-size: 11pt;
    line-height: 1.4;
    color: #333;
}

h1 {
    font-size: 24pt;
    margin-bottom: 0.5em;
    padding-bottom: 0.3em;
    border-bottom: 2px solid #333;
}

h2 {
    font-size: 14pt;
    margin-top: 1.5em;
    margin-bottom: 0.8em;
    padding-bottom: 0.2em;
    border-bottom: 1px solid #666;
}

.contact-info {
    font-size: 10pt;
    margin-bottom: 1em;
    color: #555;
}

.section-entry {
    display: flex;
    margin-bottom: 1.2em;
    page-break-inside: avoid;
}

.entry-left {
    flex: 0 0 40%;
    font-weight: bold;
    padding-right: 1em;
}

.entry-right {
    flex: 1;
}

.entry-right .position {
    font-weight: bold;
    margin-bottom: 0.3em;
}

.entry-right .details {
    font-size: 10pt;
    color: #555;
    margin-bottom: 0.5em;
}

ul {
    margin: 0.5em 0;
    padding-left: 1.5em;
}

li {
    margin-bottom: 0.3em;
}

.skills-list {
    margin-bottom: 0.8em;
}

.skills-list strong {
    display: inline-block;
    min-width: 180px;
}
"""

# CSS Template for Cover Letter (simple single-column)
COVER_LETTER_CSS_TEMPLATE = """
@page {
    size: A4;
    margin: 2.5cm;
}

body {
    font-family: 'Times New Roman', 'Georgia', serif;
    font-size: 12pt;
    line-height: 1.6;
    color: #000;
}

.header {
    margin-bottom: 2em;
    font-size: 11pt;
}

.date {
    margin-bottom: 2em;
}

.greeting {
    margin-bottom: 1.5em;
}

p {
    margin-bottom: 1.2em;
    text-align: justify;
}

.betreff {
    margin-bottom: 1.5em;
}

.closing {
    margin-top: 2em;
}

.signature {
    margin-top: 3em;
}
"""

# ============================================
# Tools (Agent Capabilities)
# ============================================

@tool
def analyze_cv_job_alignment(cv_json: str, job_json: str) -> str:
    """
    Analyze how well the CV matches job requirements and create a tailoring plan.
    
    This tool performs gap analysis between the candidate's CV and the target job,
    identifying strengths to emphasize and areas where the CV can be optimized.
    
    IMPORTANT: This tool receives pre-processed JSON from cv_agent and job_agent.
    No extraction is needed - just analysis of structured data.
    
    Args:
        cv_json: ResumeInfo JSON string from cv_agent output
        job_json: JobRequirements JSON string from job_agent output
        
    Returns:
        JSON string with CVTailoringPlan containing specific, actionable suggestions
        
    Example:
        >>> plan = analyze_cv_job_alignment(cv_data, job_data)
        >>> # Returns detailed plan with matching experiences, skills, keywords, etc.
    """
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)
    structured_llm = llm.with_structured_output(CVTailoringPlan)
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", """
            You are an expert career counselor conducting gap analysis.
            Your task is to compare the candidate's CV with job requirements and identify:

            1. MATCHING EXPERIENCES: Which work experiences or roles align with job responsibilities?
            2. MATCHING SKILLS: Which technical and soft skills overlap between CV and job?
            3. RELEVANT PROJECTS: Which projects demonstrate capabilities needed for this job?
            4. KEYWORDS: What job-specific terms should appear in the tailored CV?
            5. REORDERING: Should sections be reorganized to put most relevant items first?
            6. EMPHASIS: What achievements deserve stronger highlighting?

            CRITICAL CONSTRAINTS:
            - Only work with EXISTING content from the CV
            - Never suggest fabricating experience or skills
            - Focus on emphasizing and reframing, not inventing
            - Be specific with actionable suggestions
            Provide concrete, implementable recommendations.
        """),
        ("user", """
        Candidate's CV:
        {cv}

        Job Requirements:
        {job}

        Analyze the alignment and create a detailed tailoring plan.
        """)
    ])
    
    chain = prompt | structured_llm
    result = chain.invoke({"cv": cv_json, "job": job_json})
    
    return result.model_dump_json(indent=2)

@tool
def generate_tailored_cv_html(cv_json: str, tailoring_plan_json: str) -> str:
    """
    Generate HTML content for a tailored CV based on original CV and tailoring plan.
    
    This tool creates HTML that maintains the original two-column table layout
    while incorporating the suggestions from the tailoring plan.
    
    Args:
        cv_json: Original ResumeInfo JSON from cv_agent
        tailoring_plan_json: CVTailoringPlan JSON from analyze_cv_job_alignment
        
    Returns:
        HTML string representing the tailored CV (body content only, no <html> wrapper)
        
    Note:
        The HTML uses semantic structure with classes for styling:
        - .section-entry for each CV entry
        - .entry-left for institution/company names
        - .entry-right for details and descriptions
    """
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.3)
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", """
        You are generating HTML for a tailored CV.
        LAYOUT REQUIREMENTS:
        1. Maintain two-column structure:
        - Left column (40%): Institution/Company names in bold
        - Right column (60%): Roles, dates, locations, bullet points

        2. Structure:
        - <h1> for candidate name
        - <div class="contact-info"> for contact details
        - <h2> for each section (Education, Experience, Skills & Projects, etc.)
        - <div class="section-entry"> for each entry with:
            - <div class="entry-left"> (institution/company)
            - <div class="entry-right"> (details)

        3. Incorporate tailoring plan:
        - Reorder sections/items per suggestions
        - Emphasize relevant experiences and projects
        - Naturally weave in keywords from the plan
        - Highlight matching skills prominently

        4. Preserve authenticity:
        - Keep the candidate's voice and writing style
        - Don't fabricate content
        - Enhance descriptions with keywords, don't replace them

        RETURN: Only the HTML body content (no <html>, <head>, or <body> tags).
        Use semantic HTML with the specified class names for proper styling."""),
        ("user", """Original CV:
        {cv}

        Tailoring Plan:
        {plan}

        Generate the tailored CV HTML.""")
    ])
    
    chain = prompt | llm
    result = chain.invoke({"cv": cv_json, "plan": tailoring_plan_json})
    
    return result.content

@tool
def generate_cover_letter_content(cv_json: str, job_json: str, company_json: str = "", language: str = "english") -> str:
    """
    Generate tailored cover letter content for the job application.

    Creates a compelling, authentic cover letter that connects the candidate's
    background to the job requirements and company culture.

    Args:
        cv_json: ResumeInfo JSON from cv_agent
        job_json: JobRequirements JSON from job_agent
        company_json: Optional CompanyInfo JSON from job_agent for company-specific points
        language: Output language/format.
            "english" (default) → standard English cover letter (3-4 paragraphs)
            "german"            → formal German Anschreiben in Sie-form with
                                  Betreff, Einleitung, Hauptteil, Schlussteil,
                                  and Grußformel

    Returns:
        JSON string with CoverLetterContent structure.
        For English: opening/body/closing paragraphs populated; betreff and
        grussformel are empty strings.
        For German: betreff and grussformel populated; opening_paragraph →
        Einleitung, body_paragraph_1/2 → Hauptteil, closing_paragraph →
        Schlussteil; language field set to "german".

    Note:
        English letter structure (3-4 paragraphs):
        - Opening: Express interest and how you learned about the role
        - Body 1-2: Connect relevant experiences to job requirements
        - Body 3 (optional): Address company-specific points
        - Closing: Call to action and appreciation

        German Anschreiben structure:
        - Betreff: Subject line naming the exact position
        - Anrede: Formal salutation (generated separately as recipient_info)
        - Einleitung: Why applying, brief self-introduction
        - Hauptteil (1-2 paragraphs): Specific qualifications and experience
        - Schlussteil: Motivation summary, request for interview
        - Grußformel: Formal valediction
    """
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.4)
    structured_llm = llm.with_structured_output(CoverLetterContent)

    # Build context with optional company info
    company_context = ""
    if company_json and company_json.strip():
        company_context = f"\n\nCompany Information:\n{company_json}"

    if language.lower() == "german":
        prompt = ChatPromptTemplate.from_messages([
            ("system", """Du schreibst ein professionelles Anschreiben auf Deutsch für eine Bewerbung.

        FORMALE VORGABEN:
        1. Sprache: Ausschließlich Deutsch, formelles Sie durchgehend
        2. Register: Formell-sachlich, klar und präzise — kein Marketingjargon
        3. Länge: 250–350 Wörter Gesamttext (ohne Betreff und Grußformel)
        4. Aufbau:
           - betreff: Einzeiliger Betreff, der die exakte Stellenbezeichnung nennt,
             z. B. „Bewerbung als Senior Software Engineer (Ref.-Nr. 12345)"
           - opening_paragraph (Einleitung): Ein kurzer Absatz — warum diese Stelle,
             kurze Selbstvorstellung
           - body_paragraph_1 (1. Hauptteil): Konkrete Qualifikationen und
             Erfahrungen, die direkt auf die Stellenanforderungen einzahlen
           - body_paragraph_2 (2. Hauptteil): Weitere relevante Kompetenzen oder
             Projekte; bei Bedarf Unternehmenbezug
           - body_paragraph_3 (optional): Nur befüllen, wenn Unternehmensinformationen
             vorliegen und ein spezifischer Mehrwert hergestellt werden kann
           - closing_paragraph (Schlussteil): Motivation bekräftigen,
             Gesprächswunsch äußern, Verfügbarkeit nennen
           - grussformel: „Mit freundlichen Grüßen" (Standard) oder eine
             situationsangemessene Alternative

        5. Inhaltliche Prinzipien:
           - KONKRET: Echte Erfahrungen und Kenntnisse aus dem Lebenslauf referenzieren
           - AUSGERICHTET: Zeigen, dass die Anforderungen der Stelle verstanden wurden
           - MEHRWERT: Erklären, welchen Beitrag die Kandidatin/der Kandidat leistet
           - AUTHENTISCH: Den Stil der Bewerberin/des Bewerbers wahren

        VERMEIDEN:
        - Generische Floskeln wie „Hiermit bewerbe ich mich ..."
        - Reine Wiederholung des Lebenslaufs
        - Übertriebenes Eigenlob oder unterwürfige Formulierungen
        - Anglizismen und Buzzwords
        - Du-Form oder Kumpelton
        - Englische Wörter oder Satzteile"""),
            ("user", """Lebenslauf der Kandidatin / des Kandidaten:
        {cv}

        Stellenanforderungen:
        {job}
        {company}

        Schreibe ein überzeugendes, maßgeschneidertes Anschreiben auf Deutsch (Sie-Form).
        Setze language=\"german\" und befülle betreff sowie grussformel entsprechend.""")
        ])
    else:
        prompt = ChatPromptTemplate.from_messages([
            ("system", """You are writing a professional cover letter.

        WRITING GUIDELINES:
        1. Tone: Professional yet personable, enthusiastic but not over-the-top
        2. Length: 300-400 words total (concise and impactful)
        3. Structure:
        - Opening: Hook with genuine interest in role/company
        - Body: Connect 2-3 most relevant experiences to job needs
        - Closing: Strong call to action

        4. Content principles:
        - Be SPECIFIC: Reference actual experiences and skills from CV
        - Show ALIGNMENT: Demonstrate understanding of job requirements
        - Add VALUE: Explain what you'll bring to the role
        - Be AUTHENTIC: Maintain candidate's voice

        5. If company info provided:
        - Reference company values or culture
        - Mention relevant company news or initiatives
        - Show you've done research

        AVOID:
        - Generic statements that could apply to any job
        - Simply repeating CV content
        - Explaining why YOU want the job (focus on what YOU bring)
        - Desperation or begging tone
        - Clichés and buzzwords"""),
            ("user", """Candidate's CV:
        {cv}

        Job Requirements:
        {job}
        {company}

        Write a compelling, tailored cover letter.""")
        ])

    chain = prompt | structured_llm
    result = chain.invoke({"cv": cv_json, "job": job_json, "company": company_context})

    return result.model_dump_json(indent=2)

@tool
def generate_cv_pdf(html_content: str, output_filename: str, applicant_name: str) -> str:
    """
    Generate final CV PDF from HTML content with professional styling.
    
    WARNING: Only call this tool after the user has explicitly approved the CV content!
    
    Args:
        html_content: HTML body content from generate_tailored_cv_html
        output_filename: Filename for the PDF (e.g., "kevin_ha_cv_tailored.pdf")
        applicant_name: Full name for document title metadata
        
    Returns:
        Success message with the full path to the generated PDF file
        
    Raises:
        Exception: If PDF generation fails or output directory is not writable
        
    Example:
        >>> result = generate_cv_pdf(html, "john_doe_cv.pdf", "John Doe")
        >>> print(result)
        "CV PDF generated successfully! The file 'john_doe_cv.pdf' is ready for download."
    """
    try:
        # Import WeasyPrint (lazy loading)
        HTML, CSS = _import_weasyprint()
        
        # Ensure output directory exists
        output_dir = DEFAULT_OUTPUT_DIR
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Construct full output path
        output_path = output_dir / output_filename
        
        # Build complete HTML document
        # Escape applicant_name to prevent HTML injection
        escaped_name = html.escape(applicant_name)
        full_html = f"""<!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{escaped_name} - CV</title>
        </head>
        <body>
        {html_content}
        </body>
        </html>"""
        
        # Generate PDF with styling
        HTML(string=full_html).write_pdf(
            str(output_path),
            stylesheets=[CSS(string=CV_CSS_TEMPLATE)]
        )
        
        # Return user-friendly message without exposing server path
        # Filename is included for backend detection (will be improved with tool tracking)
        return f"CV PDF generated successfully! The file '{output_filename}' is ready for download."
        
    except RuntimeError as e:
        # WeasyPrint not available
        return f"Error: {str(e)}"
    except Exception as e:
        return f"Error generating CV PDF: {str(e)}"

@tool
def generate_cover_letter_pdf(content_json: str, output_filename: str, applicant_name: str, applicant_contact: str, recipient_info: str = "Hiring Manager") -> str:
    """
    Generate cover letter PDF from structured content.
    
    Creates a simple, professional cover letter PDF with proper business letter formatting.
    
    WARNING: Only call this tool after the user has explicitly approved the cover letter content!
    
    Args:
        content_json: CoverLetterContent JSON from generate_cover_letter_content (the JSON string output from that tool)
        output_filename: Filename for the PDF (e.g., "john_doe_cover_letter.pdf")
        applicant_name: Full name from CV data - extract from cv_data['name'] or cv_data['full_name']
        applicant_contact: Contact info formatted as "email | phone" - construct from cv_data['email'] and cv_data['phone']
        recipient_info: Who to address (default: "Hiring Manager") - can extract from job_data if specified
        
    Returns:
        Success message indicating the PDF was generated (filename included for backend detection)
        
    Example:
        >>> # Get CV data from system context
        >>> name = cv_data['name']  # or cv_data['full_name']
        >>> email = cv_data.get('email', 'email@example.com')
        >>> phone = cv_data.get('phone', '')
        >>> contact = f"{email} | {phone}" if phone else email
        >>> 
        >>> result = generate_cover_letter_pdf(
        ...     content_json=cover_letter_json_output,
        ...     output_filename="john_doe_cover_letter.pdf",
        ...     applicant_name=name,
        ...     applicant_contact=contact,
        ...     recipient_info="Hiring Manager"
        ... )
    """
    try:
        # Import WeasyPrint (lazy loading)
        HTML, CSS = _import_weasyprint()
        
        # Parse content with better error handling
        try:
            content = json.loads(content_json)
        except json.JSONDecodeError as e:
            return f"Error: Invalid JSON format for cover letter content. {str(e)}"
        
        # Validate required fields
        required_fields = ['opening_paragraph', 'body_paragraph_1', 'body_paragraph_2', 'closing_paragraph']
        missing_fields = [field for field in required_fields if field not in content]
        if missing_fields:
            return f"Error: Cover letter content is missing required fields: {', '.join(missing_fields)}"
        
        # Ensure output directory exists
        output_dir = DEFAULT_OUTPUT_DIR
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Construct full output path
        output_path = output_dir / output_filename
        
        # Determine language/format
        is_german = content.get('language', 'english').lower() == 'german'

        # Build body paragraphs list (excluding empty optional paragraph)
        paragraphs = [
            content['opening_paragraph'],
            content['body_paragraph_1'],
            content['body_paragraph_2']
        ]
        if content.get('body_paragraph_3', '').strip():
            paragraphs.append(content['body_paragraph_3'])
        paragraphs.append(content['closing_paragraph'])

        # Generate paragraph HTML with proper escaping
        paragraphs_html = '\n'.join([f'<p>{html.escape(p)}</p>' for p in paragraphs])

        # Get current date
        from datetime import datetime
        if is_german:
            # German locale date format: "18. Juni 2026"
            GERMAN_MONTHS = {
                1: "Januar", 2: "Februar", 3: "März", 4: "April",
                5: "Mai", 6: "Juni", 7: "Juli", 8: "August",
                9: "September", 10: "Oktober", 11: "November", 12: "Dezember"
            }
            now = datetime.now()
            current_date = f"{now.day}. {GERMAN_MONTHS[now.month]} {now.year}"
        else:
            current_date = datetime.now().strftime("%B %d, %Y")

        # Escape user-provided data to prevent HTML injection
        escaped_name = html.escape(applicant_name)
        escaped_contact = html.escape(applicant_contact)
        escaped_recipient = html.escape(recipient_info)

        if is_german:
            betreff = html.escape(content.get('betreff', ''))
            grussformel = html.escape(content.get('grussformel', 'Mit freundlichen Grüßen'))
            betreff_html = f'<div class="betreff"><strong>{betreff}</strong></div>' if betreff else ''
            salutation = f"Sehr geehrte/r {escaped_recipient}," if escaped_recipient and escaped_recipient != "Hiring Manager" else "Sehr geehrte Damen und Herren,"
            full_html = f"""<!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>Anschreiben - {escaped_name}</title>
        </head>
        <body>
            <div class="header">
                <strong>{escaped_name}</strong><br>
                {escaped_contact}
            </div>

            <div class="date">
                {current_date}
            </div>

            {betreff_html}

            <div class="greeting">
                {salutation}
            </div>

            {paragraphs_html}

            <div class="closing">
                {grussformel}
            </div>

            <div class="signature">
                {escaped_name}
            </div>
        </body>
        </html>"""
        else:
            full_html = f"""<!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>Cover Letter - {escaped_name}</title>
        </head>
        <body>
            <div class="header">
                <strong>{escaped_name}</strong><br>
                {escaped_contact}
            </div>

            <div class="date">
                {current_date}
            </div>

            <div class="greeting">
                Dear {escaped_recipient},
            </div>

            {paragraphs_html}

            <div class="closing">
                Sincerely,
            </div>

            <div class="signature">
                {escaped_name}
            </div>
        </body>
        </html>"""
        
        # Generate PDF with styling
        HTML(string=full_html).write_pdf(
            str(output_path),
            stylesheets=[CSS(string=COVER_LETTER_CSS_TEMPLATE)]
        )
        
        # Return user-friendly message without exposing server path
        # Filename is included for backend detection (will be improved with tool tracking)
        return f"Cover letter PDF generated successfully! The file '{output_filename}' is ready for download."
        
    except RuntimeError as e:
        # WeasyPrint not available
        return f"Error: {str(e)}"
    except Exception as e:
        return f"Error generating cover letter PDF: {str(e)}"

@tool
def generate_cv_docx(cv_json: str, output_filename: str, applicant_name: str) -> str:
    """
    Generate a Word document (.docx) with CV content in basic structure.
    
    WARNING: Only call this tool after the user has explicitly approved the CV content!
    
    This creates a simple Word document with structured content. The formatting is minimal
    so users can easily customize it in Microsoft Word or Google Docs to match their preferences.
    
    Args:
        cv_json: ResumeInfo JSON string from cv_agent (the original CV data)
        output_filename: Filename for the Word document (e.g., "kevin_ha_cv_tailored.docx")
        applicant_name: Full name for document title
    
    Returns:
        Success message with the filename for download
    
    Example:
        >>> result = generate_cv_docx(cv_data, "john_doe_cv.docx", "John Doe")
        >>> print(result)
    """
    try:
        from docx import Document
        from docx.shared import Pt
        
        # Parse CV data
        cv_data = json.loads(cv_json)
        
        # Ensure output directory exists
        output_dir = DEFAULT_OUTPUT_DIR
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Construct full output path
        output_path = output_dir / output_filename
        
        # Create document
        doc = Document()
        
        # Title (Name) - Level 1 heading
        title = doc.add_heading(cv_data.get('name', applicant_name), level=1)
        
        # Contact information
        contact_parts = []
        if cv_data.get('email'):
            contact_parts.append(cv_data['email'])
        if cv_data.get('phone'):
            contact_parts.append(cv_data['phone'])
        if cv_data.get('location'):
            contact_parts.append(cv_data['location'])
        if cv_data.get('linkedin_url'):
            contact_parts.append(f"LinkedIn: {cv_data['linkedin_url']}")
        if cv_data.get('github_url'):
            contact_parts.append(f"GitHub: {cv_data['github_url']}")
        if cv_data.get('portfolio_url'):
            contact_parts.append(f"Portfolio: {cv_data['portfolio_url']}")
        
        if contact_parts:
            contact_para = doc.add_paragraph(' | '.join(contact_parts))
            contact_para.style.font.size = Pt(10)
        
        # Education section
        if cv_data.get('education'):
            doc.add_heading('Education', level=2)
            for edu in cv_data['education']:
                edu_para = doc.add_paragraph()
                if edu.get('degree'):
                    edu_para.add_run(edu['degree']).bold = True
                if edu.get('institution'):
                    if edu.get('degree'):
                        edu_para.add_run(f" - {edu['institution']}")
                    else:
                        edu_para.add_run(edu['institution']).bold = True
                if edu.get('dates'):
                    edu_para.add_run(f" ({edu['dates']})")
                if edu.get('location'):
                    edu_para.add_run(f", {edu['location']}")
                if edu.get('gpa'):
                    edu_para.add_run(f" - GPA: {edu['gpa']}")
        
        # Experience section
        if cv_data.get('experience'):
            doc.add_heading('Experience', level=2)
            for exp in cv_data['experience']:
                # Position and company
                exp_para = doc.add_paragraph()
                if exp.get('position'):
                    exp_para.add_run(exp['position']).bold = True
                if exp.get('company'):
                    if exp.get('position'):
                        exp_para.add_run(f" at {exp['company']}")
                    else:
                        exp_para.add_run(exp['company']).bold = True
                if exp.get('dates'):
                    exp_para.add_run(f" ({exp['dates']})")
                if exp.get('location'):
                    exp_para.add_run(f" - {exp['location']}")
                
                # Responsibilities/bullet points
                if exp.get('responsibilities'):
                    for resp in exp['responsibilities']:
                        doc.add_paragraph(resp, style='List Bullet')
        
        # Skills section
        if cv_data.get('skills'):
            doc.add_heading('Skills', level=2)
            skills_text = ', '.join(cv_data['skills'])
            doc.add_paragraph(skills_text)
        
        # Projects section
        if cv_data.get('projects'):
            doc.add_heading('Projects', level=2)
            for project in cv_data['projects']:
                proj_para = doc.add_paragraph()
                if project.get('name'):
                    proj_para.add_run(project['name']).bold = True
                if project.get('description'):
                    if project.get('name'):
                        proj_para.add_run(f": {project['description']}")
                    else:
                        proj_para.add_run(project['description'])
                if project.get('technologies'):
                    tech_text = ', '.join(project['technologies']) if isinstance(project['technologies'], list) else project['technologies']
                    doc.add_paragraph(f"Technologies: {tech_text}", style='List Bullet')
        
        # Leadership & Activities section
        if cv_data.get('leadership_activities'):
            doc.add_heading('Leadership & Activities', level=2)
            for activity in cv_data['leadership_activities']:
                act_para = doc.add_paragraph()
                if activity.get('role'):
                    act_para.add_run(activity['role']).bold = True
                if activity.get('organization'):
                    if activity.get('role'):
                        act_para.add_run(f" - {activity['organization']}")
                    else:
                        act_para.add_run(activity['organization']).bold = True
                if activity.get('dates'):
                    act_para.add_run(f" ({activity['dates']})")
                if activity.get('description'):
                    doc.add_paragraph(activity['description'], style='List Bullet')
        
        # Save document
        doc.save(str(output_path))
        
        return f"CV Word document generated successfully! The file '{output_filename}' is ready for download. You can customize the formatting in Microsoft Word or Google Docs."
        
    except ImportError:
        return "Error: python-docx library is not installed. Please install it to generate Word documents."
    except json.JSONDecodeError as e:
        return f"Error: Invalid CV JSON data: {str(e)}"
    except Exception as e:
        return f"Error generating CV Word document: {str(e)}"

@tool
def generate_cover_letter_docx(content_json: str, output_filename: str, applicant_name: str, applicant_contact: str, recipient_info: str = "Hiring Manager") -> str:
    """
    Generate a Word document (.docx) with cover letter content in basic structure.
    
    WARNING: Only call this tool after the user has explicitly approved the cover letter content!
    
    This creates a simple Word document with structured cover letter content. The formatting is minimal
    so users can easily customize it in Microsoft Word or Google Docs.
    
    Args:
        content_json: CoverLetterContent JSON from generate_cover_letter_content (the JSON string output)
        output_filename: Filename for the Word document (e.g., "john_doe_cover_letter.docx")
        applicant_name: Full name from CV data
        applicant_contact: Contact info formatted as "email | phone"
        recipient_info: Who to address (default: "Hiring Manager")
    
    Returns:
        Success message with the filename for download
    
    Example:
        >>> result = generate_cover_letter_docx(
        ...     content_json=cover_letter_json_output,
        ...     output_filename="john_doe_cover_letter.docx",
        ...     applicant_name="John Doe",
        ...     applicant_contact="john@email.com | (123) 456-7890)",
        ...     recipient_info="Hiring Manager"
        ... )
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime

        # Parse cover letter content
        content_data = json.loads(content_json)

        is_german = content_data.get('language', 'english').lower() == 'german'

        # Ensure output directory exists
        output_dir = DEFAULT_OUTPUT_DIR
        output_dir.mkdir(parents=True, exist_ok=True)

        # Construct full output path
        output_path = output_dir / output_filename

        # Create document
        doc = Document()

        # Header with name and contact
        header_para = doc.add_paragraph()
        header_para.add_run(applicant_name).bold = True
        header_para.add_run(f"\n{applicant_contact}")

        # Date
        if is_german:
            GERMAN_MONTHS = {
                1: "Januar", 2: "Februar", 3: "März", 4: "April",
                5: "Mai", 6: "Juni", 7: "Juli", 8: "August",
                9: "September", 10: "Oktober", 11: "November", 12: "Dezember"
            }
            now = datetime.now()
            current_date = f"{now.day}. {GERMAN_MONTHS[now.month]} {now.year}"
        else:
            current_date = datetime.now().strftime("%B %d, %Y")

        date_para = doc.add_paragraph(current_date)
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # German Betreff (subject line) before salutation
        if is_german:
            betreff = content_data.get('betreff', '').strip()
            if betreff:
                betreff_para = doc.add_paragraph()
                betreff_para.add_run(betreff).bold = True

        # Greeting / Salutation
        greeting_para = doc.add_paragraph()
        if is_german:
            if recipient_info and recipient_info != "Hiring Manager":
                salutation = f"Sehr geehrte/r {recipient_info},"
            else:
                salutation = "Sehr geehrte Damen und Herren,"
        else:
            salutation = f"Dear {recipient_info},"
        greeting_para.add_run(salutation)

        # Body paragraphs — read from CoverLetterContent fields
        body_fields = [
            content_data.get('opening_paragraph', ''),
            content_data.get('body_paragraph_1', ''),
            content_data.get('body_paragraph_2', ''),
        ]
        body_para_3 = content_data.get('body_paragraph_3', '').strip()
        if body_para_3:
            body_fields.append(body_para_3)
        body_fields.append(content_data.get('closing_paragraph', ''))

        for para_text in body_fields:
            if para_text.strip():
                doc.add_paragraph(para_text)

        # Closing / Grußformel
        if is_german:
            grussformel = content_data.get('grussformel', 'Mit freundlichen Grüßen')
            closing_para = doc.add_paragraph(grussformel)
        else:
            closing_para = doc.add_paragraph("Sincerely,")
        closing_para.space_after = Pt(12)

        # Signature
        signature_para = doc.add_paragraph(applicant_name)
        signature_para.space_before = Pt(24)

        # Save document
        doc.save(str(output_path))

        return f"Cover letter Word document generated successfully! The file '{output_filename}' is ready for download. You can customize the formatting in Microsoft Word or Google Docs."

    except ImportError:
        return "Error: python-docx library is not installed. Please install it to generate Word documents."
    except json.JSONDecodeError as e:
        return f"Error: Invalid cover letter JSON data: {str(e)}"
    except Exception as e:
        return f"Error generating cover letter Word document: {str(e)}"

# ============================================
# Enhanced Tailoring Tools (Phase 1 & 2)
# ============================================

@tool
def generate_job_summary(job_json: str) -> str:
    """
    Generate a brief summary of the job posting including role, responsibilities, and required skills.
    
    Args:
        job_json: JobRequirements JSON string from job_agent output
        
    Returns:
        JSON string with job summary containing:
        - role: Job title and level
        - responsibilities: Key responsibilities
        - required_skills: List of required skills
        - preferred_skills: Nice-to-have skills
        - key_notes: Other important points
    """
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)
    structured_llm = llm.with_structured_output(JobSummary)

    prompt = ChatPromptTemplate.from_messages([
        ("system", (
            "You are summarizing a job posting. "
            "Extract the role title with seniority level, all key responsibilities, "
            "required skills, preferred/nice-to-have skills, and any other important "
            "notes (location, employment type, culture fit signals, etc.)."
        )),
        ("user", "Job posting data:\n{job}")
    ])

    try:
        chain = prompt | structured_llm
        summary: JobSummary = chain.invoke({"job": job_json})
        return summary.model_dump_json()
    except Exception as e:
        return json.dumps({
            "role": "Unknown",
            "responsibilities": [],
            "required_skills": [],
            "preferred_skills": [],
            "key_notes": [f"Summarization failed: {str(e)[:120]}"]
        })

def _extract_compatibility_report(compatibility_data: Dict[str, Any]) -> Optional[CompatibilityReport]:
    """
    Accept either a v2 CompatibilityReport directly or the legacy shim shape
    (with `report_v2` embedded). Returns None if the input is malformed.
    """
    if not isinstance(compatibility_data, dict):
        return None
    candidate = compatibility_data
    if "report_v2" in compatibility_data and isinstance(compatibility_data["report_v2"], dict):
        candidate = compatibility_data["report_v2"]
    try:
        if "aggregate_score" in candidate or "dimensions" in candidate:
            return CompatibilityReport.model_validate(candidate)
    except Exception:
        return None
    return None


def _intensity_for_level(level: str) -> str:
    return {
        "excellent": "minor",
        "high": "minor",
        "medium": "moderate",
        "low": "major",
        "unknown": "moderate",
    }.get(level, "moderate")


def _build_directives_from_report(report: CompatibilityReport) -> Dict[str, Any]:
    """
    Translate a CompatibilityReport into concrete per-section directives the
    downstream tailoring tools can act on without needing further LLM reasoning.
    """
    gap = report.gap_analysis
    directives: List[Dict[str, Any]] = []
    focus_areas: List[str] = []

    # Hard skills: emit per-transferable directives with bridge bullets
    for sk in gap.transferable_skills:
        directives.append({
            "type": "inject_bridge_bullet",
            "section": "experience_or_projects",
            "required_skill": sk.required_skill,
            "candidate_skill": sk.matched_with,
            "kind": sk.kind,
            "transferability": sk.transferability,
            "bullet": sk.bridge_bullet or (
                f"Leveraged {sk.matched_with} on prior work; "
                f"{sk.matched_with} and {sk.required_skill} share core paradigms."
                if sk.matched_with else f"Highlight transferable experience related to {sk.required_skill}."
            ),
            "rationale": sk.rationale,
        })

    if gap.matched_skills:
        focus_areas.append("emphasize_matched_skills")
        directives.append({
            "type": "emphasize_skills_section",
            "skills_to_surface": [m.matched_with or m.required_skill for m in gap.matched_skills],
            "rationale": "These job-required skills are explicitly held by the candidate; surface them prominently.",
        })

    if gap.transferable_skills:
        focus_areas.append("inject_bridge_bullets")

    if gap.missing_skills:
        focus_areas.append("address_missing_skills")
        omit_count = sum(1 for m in gap.missing_skills if m.transferability < 0.2)
        soften_count = len(gap.missing_skills) - omit_count
        directives.append({
            "type": "handle_missing_skills",
            "missing": [m.required_skill for m in gap.missing_skills],
            "guidance": (
                f"Omit unfamiliar tooling that has no transferable foothold "
                f"({omit_count} skills); for the remaining {soften_count}, "
                "use neutral phrasing that emphasizes adjacent experience without overclaiming."
            ),
        })

    # Per-dimension directives
    dim_lookup = {d.name: d for d in report.dimensions}
    if dim_lookup.get("experience") and dim_lookup["experience"].score < 0.5:
        focus_areas.append("amplify_experience_outcomes")
        directives.append({
            "type": "amplify_outcomes",
            "section": "experience",
            "guidance": "Surface quantified outcomes (scope, scale, impact) in the most relevant roles to compensate for the years gap.",
        })
    if dim_lookup.get("seniority") and dim_lookup["seniority"].score < 0.5:
        focus_areas.append("strengthen_seniority_signals")
        directives.append({
            "type": "strengthen_seniority_language",
            "section": "experience",
            "guidance": "Where truthful, use scope/leadership verbs (led, owned, mentored, drove) and add team/scope sizing.",
        })
    if dim_lookup.get("ats_keywords") and dim_lookup["ats_keywords"].score < 0.55:
        focus_areas.append("inject_ats_keywords")
        directives.append({
            "type": "inject_keywords",
            "section": "skills_and_summary",
            "guidance": "Insert exact-match required-skill terms in the skills section and summary so ATS keyword filters pass.",
        })
    if dim_lookup.get("domain") and dim_lookup["domain"].score < 0.5:
        focus_areas.append("reframe_domain_language")
        directives.append({
            "type": "reframe_domain",
            "section": "summary_and_titles",
            "guidance": "Reframe role titles and summary to use the job's domain terminology where the candidate has matching experience.",
        })

    if gap.over_qualified_signals:
        focus_areas.append("temper_over_qualification")
        directives.append({
            "type": "temper_seniority",
            "guidance": "Keep senior-only details concise; emphasize hands-on contribution to avoid signaling over-qualification.",
            "signals": gap.over_qualified_signals,
        })

    return {
        "scoring_version": "v2",
        "aggregate_score": report.aggregate_score,
        "level": report.level,
        "intensity": _intensity_for_level(report.level),
        "strategy": report.interpretation,
        "focus_areas": list(dict.fromkeys(focus_areas)),  # de-duped, stable order
        "directives": directives,
        "summary": {
            "matched_skills_count": len(gap.matched_skills),
            "transferable_skills_count": len(gap.transferable_skills),
            "missing_skills_count": len(gap.missing_skills),
            "over_qualified_signals_count": len(gap.over_qualified_signals),
        },
    }


@tool
def decide_tailoring_strategy(compatibility_score_json: str, cv_json: str, job_json: str) -> str:
    """
    Decide tailoring strategy from a CompatibilityReport (v2) or the legacy
    compatibility score response.

    Produces concrete per-skill directives:
    - inject_bridge_bullet for each transferable skill (with draft bullet)
    - emphasize_skills_section listing directly matched skills
    - handle_missing_skills with omit/soften guidance
    - amplify_outcomes / strengthen_seniority_language / inject_keywords /
      reframe_domain when those dimensions score low
    - temper_seniority if over-qualification signals were detected

    These directives are consumed directly by select_prioritize_content and
    rewrite_enhance_content - no further LLM reasoning required to pick a
    strategy.

    Args:
        compatibility_score_json: JSON from calculate_compatibility_score (v1
            shim) or calculate_compatibility_score_v2.
        cv_json: ResumeInfo JSON string.
        job_json: JobRequirements JSON string.

    Returns:
        JSON string with strategy, intensity, focus_areas, and a structured
        directives list.
    """
    try:
        compatibility_data = json.loads(compatibility_score_json) if isinstance(compatibility_score_json, str) else compatibility_score_json
    except json.JSONDecodeError as e:
        return json.dumps({
            "error": f"Invalid compatibility JSON: {str(e)}",
            "strategy": "standard",
            "intensity": "moderate",
            "focus_areas": [],
            "directives": [],
        })

    report = _extract_compatibility_report(compatibility_data)

    # If we couldn't parse a v2 report (e.g. caller passed a raw v1 dict from
    # a non-shim source), recompute v2 on the fly so we always have a rich
    # structure to act on.
    if report is None:
        try:
            cv_data = json.loads(cv_json) if isinstance(cv_json, str) else cv_json
            job_data = json.loads(job_json) if isinstance(job_json, str) else job_json
            if isinstance(cv_data, dict) and isinstance(job_data, dict):
                report = _calculate_compatibility_score_v2_internal(cv_data, job_data)
        except Exception:
            report = None

    if report is None:
        # Last-resort fallback to the old level-based defaults
        score = 0.0
        level = "unknown"
        if isinstance(compatibility_data, dict):
            score = float(compatibility_data.get("compatibility_score", 0.0) or 0.0)
            level = str(compatibility_data.get("level", "unknown"))
        fallback = {
            "scoring_version": "fallback",
            "aggregate_score": score,
            "level": level,
            "intensity": _intensity_for_level(level),
            "strategy": "Could not parse compatibility report; using level-based default.",
            "focus_areas": ["skills", "keywords", "section_ordering"],
            "directives": [],
        }
        return json.dumps(fallback)

    return json.dumps(_build_directives_from_report(report), indent=2)

@tool
def select_prioritize_content(cv_json: str, job_json: str, tailoring_strategy_json: str) -> str:
    """
    Select and prioritize the most relevant resume content based on job requirements.

    Args:
        cv_json: ResumeInfo JSON string
        job_json: JobRequirements JSON string
        tailoring_strategy_json: JSON string from decide_tailoring_strategy

    Returns:
        JSON string with selected bullets, recommended section order, and emphasis plan
        (fields: selected_bullets, section_order, sections_to_emphasize, items_to_de_emphasize)
    """
    try:
        strategy_data = json.loads(tailoring_strategy_json)
        focus_areas: List[str] = strategy_data.get("focus_areas", [])
        directives: List[str] = strategy_data.get("directives", [])

        llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)
        structured_llm = llm.with_structured_output(SelectedContent)

        prompt = ChatPromptTemplate.from_messages([
            ("system", (
                "You are a career consultant selecting and prioritising resume content for a specific job. "
                "Your decisions must be driven by the tailoring strategy directives provided. "
                "CRITICAL RULES:\n"
                "- Only work with content that EXISTS in the CV — do not invent bullets.\n"
                "- Score each bullet 0–1 for relevance to the job requirements.\n"
                "- Recommend a section ordering that puts the most impactful sections first.\n"
                "- Mark low-value items for de-emphasis (do not delete — the candidate reviews first)."
            )),
            ("user", (
                "CV Data:\n{cv}\n\n"
                "Job Data:\n{job}\n\n"
                "Focus Areas: {focus}\n\n"
                "Strategy Directives:\n{directives}\n\n"
                "Select and prioritise the most relevant content."
            ))
        ])

        chain = prompt | structured_llm
        result: SelectedContent = chain.invoke({
            "cv": cv_json,
            "job": job_json,
            "focus": ", ".join(focus_areas),
            "directives": "\n".join(directives) if directives else "None provided."
        })
        return result.model_dump_json()

    except Exception as e:
        return json.dumps({
            "error": str(e),
            "selected_bullets": [],
            "section_order": [],
            "sections_to_emphasize": [],
            "items_to_de_emphasize": []
        })

@tool
def rewrite_enhance_content(cv_json: str, job_json: str, selected_content_json: str, tailoring_strategy_json: str) -> str:
    """
    Rewrite and enhance resume content to emphasize job relevance.

    Args:
        cv_json: Original ResumeInfo JSON string
        job_json: JobRequirements JSON string
        selected_content_json: JSON string from select_prioritize_content
        tailoring_strategy_json: JSON string from decide_tailoring_strategy

    Returns:
        JSON string with rewritten bullets, updated summary, and per-bullet confidence scores
        (fields: rewritten_bullets[{original, rewritten, confidence, keywords_added}],
         updated_summary, keywords_inserted)
    """
    try:
        strategy_data = json.loads(tailoring_strategy_json)
        directives: List[str] = strategy_data.get("directives", [])

        llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.3)
        structured_llm = llm.with_structured_output(RewrittenContent)

        prompt = ChatPromptTemplate.from_messages([
            ("system", (
                "You are a senior career coach rewriting resume bullets to maximise job relevance. "
                "Apply the strategy directives precisely.\n\n"
                "ABSOLUTE RULES:\n"
                "- Reword EXISTING content only — never fabricate experience, metrics, or skills.\n"
                "- Weave in job keywords naturally — forced keyword stuffing degrades ATS scores.\n"
                "- Lead each bullet with a strong action verb.\n"
                "- Keep the candidate's authentic voice.\n"
                "- Only enhance descriptions — do not replace their substance.\n"
                "- For each bullet, report a confidence score (0–1) for how well the rewrite "
                "matches the job, and list keywords actually added."
            )),
            ("user", (
                "Original CV:\n{cv}\n\n"
                "Job Requirements:\n{job}\n\n"
                "Selected Content:\n{selected}\n\n"
                "Strategy Directives:\n{directives}\n\n"
                "Rewrite and enhance the content."
            ))
        ])

        chain = prompt | structured_llm
        result: RewrittenContent = chain.invoke({
            "cv": cv_json,
            "job": job_json,
            "selected": selected_content_json,
            "directives": "\n".join(directives) if directives else "None provided."
        })
        return result.model_dump_json()

    except Exception as e:
        return json.dumps({
            "error": str(e),
            "rewritten_bullets": [],
            "updated_summary": "",
            "keywords_inserted": []
        })

# ============================================
# Agent Configuration
# ============================================

# Initialize LLM for the agent
llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.2)

# Comprehensive system prompt defining agent behavior
system_prompt = """You are a professional CV and cover letter writer specializing in tailoring application materials to specific job opportunities.

IMPORTANT CONTEXT:
- You receive PRE-PROCESSED, STRUCTURED DATA from upstream agents:
  * CV data (ResumeInfo JSON) - extracted by cv_agent
  * Job requirements (JobRequirements JSON) - extracted by job_agent  
  * Company info (CompanyInfo JSON, optional) - researched by job_agent
- Your role is to TRANSFORM this data into tailored, professional documents
- You DO NOT extract data from raw files - that's already done
- The complete CV and job data are provided in a system message at the start of our conversation
- This data persists throughout our entire conversation - always use the FULL JSON data from the system context when calling tools
- Do NOT use truncated or partial data from earlier messages - always reference the complete data from the system context

YOUR WORKFLOW (FOLLOW THIS SEQUENCE):

ENHANCED TAILORING MODE (when both CV and job data are provided):
Use this comprehensive 6-step workflow for optimal results:

STEP 1: INPUT & SUMMARIES
   - Display the extracted CV information to the user (show what was extracted, no summary needed)
   - Call generate_job_summary to create a brief summary of the job posting
   - Present both: the CV data overview and the job summary (role, responsibilities, required skills)
   - FORMATTING: You may use ## or ### headers, NO code blocks, NO lists with dashes (-)
   - Use plain text with line breaks for lists
   - Example format:
     "Overview of Resume and Job Match
     
     Job Title: [title]
     
     Responsibilities:
     [responsibility 1]
     [responsibility 2]
     
     Required Skills:
     Technical Skills: [list as plain text, separated by commas or line breaks]
     Soft Skills: [list as plain text, separated by commas or line breaks]"
   - This helps verify extraction accuracy and sets context

STEP 2: COMPATIBILITY ANALYSIS
   - Call calculate_compatibility_score with CV and job data
   - This runs the v2 multi-dimensional scorer which measures five weighted dimensions:
     * Hard Skills match (40%): per-skill transferability via direct/family/embedding/LLM cascade
     * Experience alignment (15%): years and depth of relevant experience
     * Seniority fit (15%): level match between candidate and role
     * Domain overlap (15%): industry/domain semantic similarity
     * ATS keyword density (15%): exact-keyword coverage for ATS filters
   - The tool also returns a gap analysis bucketed into: matched skills, transferable skills (with bridge bullets), and missing skills
   - Present the aggregate compatibility score, level, per-dimension breakdown, and key gaps
   - FORMATTING: You may use ## or ### headers, NO code blocks, NO lists with dashes
   - Use plain text with line breaks, colon format for labels
   - Example format:
     "Compatibility Analysis

     Compatibility Score: 0.72 (HIGH)

     Dimension Breakdown:
     Hard Skills (40%): 0.80 — strong direct and family matches
     Experience (15%): 0.70 — relevant depth present
     Seniority (15%): 0.75 — level aligns well
     Domain (15%): 0.65 — adjacent domain, some reframing needed
     ATS Keywords (15%): 0.60 — moderate keyword density

     Key Gaps:
     Transferable: Vue → React (bridge: 'Built component-driven UIs with Vue; transferable to React')
     Missing: Kubernetes (no equivalent found)

     Interpretation: Strong match. Focus on adding bridge bullets for transferable skills and injecting missing keywords."
   - Interpretation guidelines:
     * LOW (< 0.5): Many required skills/experience missing — flag before proceeding
     * MEDIUM (0.5 - 0.7): Good match with clear areas to improve
     * HIGH (> 0.7): Strong match, fine-tune and optimise
     * EXCELLENT (> 0.85): Near-perfect match, polish only
   - ERROR HANDLING: If score is 0.0 or error occurs:
     * Check if CV or job data is missing/empty
     * Try calling the tool again with valid data
     * If still failing, explain: "The compatibility calculation encountered an issue. This may be due to missing or incomplete data. Let's proceed with manual analysis instead."
     * Then use analyze_cv_job_alignment as fallback
   - AFTER presenting the score and dimension breakdown, always add a plain-language "Weak Points" section.
     This section MUST appear before Step 3 and MUST be written in plain, jargon-free sentences — no scores,
     no percentages, no technical labels. Identify the 2–4 most significant gaps and state them directly.
     Rules:
     * For each missing skill: one sentence naming the skill and why it matters for this role.
     * For each low-scoring dimension (< 0.6): one sentence explaining what the gap means in practice.
     * For transferable skills: one sentence noting the gap and the bridge that exists.
     * Do NOT pad this section — if there are no meaningful weak points, say so in one sentence.
   - Example Weak Points format:
     "Weak Points

     You have no listed experience with Kubernetes, which is a hard requirement for container
     orchestration in this role — this is the most significant gap.

     Your background is primarily in e-commerce; the role targets fintech, so some reframing of
     domain experience will be needed.

     You have not worked at the senior/lead level before; the role expects team leadership, which
     is not evidenced in your CV.

     Your Vue.js experience is transferable to React, but you will need to surface that connection
     explicitly — a bridge bullet in your experience section will help."

STEP 3: DECIDE TAILORING STRATEGY
   - Call decide_tailoring_strategy with compatibility score, CV, and job data
   - This determines:
     * Strategy: What approach to take (highlight transferable skills, reorder sections, fine-tune)
     * Intensity: How much editing needed (major/moderate/minor)
     * Focus areas: Which sections to prioritize (skills, projects, section_ordering, keywords, etc.)
   - BEFORE presenting the strategy details, open with a brief bridging paragraph that references
     the weak points identified in Step 2 and states explicitly how the strategy will address each one.
     Keep this to 2–4 sentences. Example:
     "Based on the weak points above, here is how the tailoring strategy will address them:
     The missing Kubernetes experience will be handled by injecting the keyword in your skills section
     and noting adjacent container experience in your project bullets. The fintech domain gap will be
     addressed by reframing your payments and fraud-detection work. The leadership gap cannot be
     fabricated — we will surface any mentoring or cross-functional coordination already in your CV."
   - Then present the full strategy
   - FORMATTING: You may use ## or ### headers, NO code blocks
   - Example format:
     "Tailoring Strategy

     Strategy: [strategy description]
     Intensity: [major/moderate/minor]
     Focus Areas: [list as plain text, separated by commas]"
   - If compatibility is LOW, inform user and ask if they want to proceed (may need to highlight transferable skills)

STEP 4: SELECT & PRIORITIZE CONTENT
   - Call select_prioritize_content with CV, job, and tailoring strategy
   - This identifies:
     * Top N bullets per section based on relevance
     * New ordering for sections and items
     * Which sections to emphasize
     * Items to de-emphasize or remove (if any)
   - Present the selected content and proposed ordering
   - FORMATTING: You may use ## or ### headers, NO code blocks, NO lists with dashes
   - Use plain text with line breaks, numbered items as "1. " format
   - WAIT for user approval before proceeding

STEP 5: REWRITE & ENHANCE CONTENT
   - Call rewrite_enhance_content with CV, job, selected content, and strategy
   - This creates:
     * Rewritten bullets with original and new versions
     * Enhanced summary/headline (if applicable)
     * Confidence scores per bullet (0-1) showing job match quality
     * List of job keywords naturally incorporated
   - Present the rewritten content with confidence scores
   - FORMATTING: You may use ## or ### headers, NO code blocks
   - Show what changed and why in clear, readable format
   - Example format:
     "Rewritten Content
     
     Original: [bullet text]
     Enhanced: [rewritten bullet text]
     Confidence: 0.85 (high match)
     
     Changes: [explanation of what changed and why]"
   - WAIT for user approval

STEP 6: USER REVIEW & EXPORT
   - Merge the rewritten content from Step 5 back into the CV JSON structure:
     * Update experience bullets with rewritten versions
     * Update summary/headline if enhanced
     * Apply new ordering from Step 4
   - Call generate_tailored_cv_html with the updated CV JSON and a tailoring plan (can create a simple plan from the strategy)
   - Present the complete tailored resume HTML for review
   - Show what changed from the original
   - Ask user if they're satisfied
   - If approved, proceed to document generation:
     * Ask user: "Ready to generate the document? I can create a PDF, Word document (.docx), or both formats. Which would you prefer?"
     * Call generate_cv_pdf and/or generate_cv_docx with the HTML content
   - If changes needed, incorporate feedback and repeat Step 5

STANDARD WORKFLOW (alternative, simpler approach):

1. ANALYSIS PHASE:
   - When user provides CV and job data, call analyze_cv_job_alignment
   - This generates a tailoring plan showing how to optimize the CV

2. REVIEW PHASE:
   - PRESENT the tailoring plan to the user in clear, readable format
   - IMPORTANT: Avoid complex markdown structure (no ### or deeper headers, no - lists, no code blocks)
   - You may use simple formatting: ## for section headers, **bold** for emphasis, and *italic* for subtle emphasis
   - Use simple line breaks, indentation, and numbered/bullet points written as plain text
   - Explain the key recommendations
   - WAIT for user feedback before proceeding
   - User might say:
     * "looks good" or "approved" → proceed to generation
     * "emphasize X more" → note the adjustment
     * "don't mention Y" → note the constraint
     * "redo the analysis" → call the tool again with updated context
   - After refining or discussing each section (Education, Experience, Skills, Projects, etc.), ask if they want to generate the document for that section
   - You can refine multiple sections before generating, but always offer generation after showing each refined section

3. CV GENERATION PHASE:
   - After approval, call generate_tailored_cv_html
   - SHOW a preview/summary of the generated HTML to the user
   - After showing each section or the complete CV, ALWAYS ask: "Ready to generate the document? I can create a PDF, Word document (.docx), or both formats. Which would you prefer?"
   - If user requests "both", "pdf and word", "both formats", or similar, generate BOTH formats
   - If user requests "pdf" or "PDF", generate only PDF
   - If user requests "word", "docx", or "Word", generate only Word document
   - If user doesn't specify, ask for clarification
   - ONLY call generate_cv_pdf and/or generate_cv_docx after EXPLICIT approval
   - Use appropriate filename (e.g., "firstname_lastname_cv_tailored.pdf" or ".docx")
   - When generating both formats, call BOTH tools and include BOTH success messages in your response
   - Note: Word documents have minimal formatting so users can easily customize them

4. COVER LETTER PHASE (if requested):
   - Before calling generate_cover_letter_content, ask the user which language/format they want:
     "Would you like the cover letter in English (standard cover letter) or German (formal Anschreiben in Sie-form)?"
     * If the user says "English" or doesn't specify: pass language="english"
     * If the user says "German", "Deutsch", "Anschreiben", or similar: pass language="german"
   - Call generate_cover_letter_content with CV, job, company data, and the chosen language parameter
   - SHOW the cover letter content to the user
     * For German: show Betreff, Anrede, and all paragraphs including Grußformel
     * For English: show opening/body/closing paragraphs as before
   - Ask for feedback or approval
   - After showing the cover letter, ALWAYS ask: "Ready to generate the document? I can create a PDF, Word document (.docx), or both formats. Which would you prefer?"
   - If user requests "both", "pdf and word", "both formats", or similar, generate BOTH formats
   - If user requests "pdf" or "PDF", generate only PDF
   - If user requests "word", "docx", or "Word", generate only Word document
   - If user doesn't specify, ask for clarification
   - ONLY call generate_cover_letter_pdf and/or generate_cover_letter_docx after EXPLICIT approval
   - When calling generate_cover_letter_pdf or generate_cover_letter_docx, you MUST provide:
     * content_json: The output from generate_cover_letter_content
     * output_filename: e.g., "firstname_lastname_cover_letter.pdf" or "firstname_lastname_anschreiben.pdf" for German
     * applicant_name: Extract from CV data (e.g., cv_data['name'])
     * applicant_contact: Format as "email | phone" from CV data (e.g., "john@email.com | (123) 456-7890")
     * recipient_info: Extract from job data or use "Hiring Manager" (for German, this will be used in the Anrede)
   - When generating both formats, call BOTH tools and include BOTH success messages in your response

CRITICAL PRINCIPLES:
ALWAYS show changes before generating PDFs
WAIT for explicit user approval at each checkpoint
Preserve the candidate's authentic voice and style
Only emphasize/refine EXISTING content - never fabricate
Be transparent about what you're changing and why
Respect user constraints and preferences
If uncertain, ask clarifying questions

NEVER generate PDFs without approval
NEVER fabricate experience, skills, or achievements
NEVER proceed to next phase without user confirmation
NEVER ignore user feedback or constraints

WHEN TO USE ENHANCED VS STANDARD WORKFLOW:
- Use ENHANCED TAILORING MODE when both CV and job data are provided and user wants comprehensive analysis
- Use STANDARD WORKFLOW for simpler, faster tailoring or when user prefers a quicker process
- You can offer both options: "I can use the enhanced tailoring mode for detailed analysis, or the standard workflow for faster results. Which would you prefer?"
- Enhanced mode provides: compatibility scores, semantic analysis, detailed strategy, confidence scores
- Standard mode provides: quick gap analysis and tailoring plan

ENHANCED TOOLS REFERENCE:
- generate_job_summary: Creates structured job summary (role, responsibilities, required/preferred skills) — Step 1
- calculate_compatibility_score: Multi-dimensional v2 scorer — 5 weighted dimensions (hard_skills 40%, experience/seniority/domain/ats_keywords 15% each) plus skill-level gap analysis — Step 2
  NOTE: calculate_semantic_similarity and calculate_bm25_score are sub-components used internally; call calculate_compatibility_score directly for Step 2 to get the full report.
- decide_tailoring_strategy: Consumes the v2 CompatibilityReport and emits per-skill directives (inject_bridge_bullet, emphasize_skills_section, handle_missing_skills, amplify_outcomes, inject_keywords, reframe_domain, temper_seniority) — Step 3
- select_prioritize_content: Selects top bullets per section with relevance scores, recommends section ordering and emphasis — Step 4
- rewrite_enhance_content: Rewrites selected bullets with action verbs and job keywords; returns per-bullet confidence scores — Step 5
- generate_cover_letter_content: Generates structured cover letter content. Accepts an optional language parameter:
  language="english" (default) → standard English cover letter (opening/body/closing paragraphs)
  language="german" → formal German Anschreiben in Sie-form (betreff, Einleitung, Hauptteil, Schlussteil, Grußformel)
  Always ask the user which language they want before calling this tool.

COMPATIBILITY THRESHOLDS:
- LOW: < 0.5 - Many gaps, consider transferable skills
- MEDIUM: 0.5 - 0.7 - Good match, optimize ordering and keywords
- HIGH: > 0.7 - Strong match, fine-tune and optimize

INTERACTION STYLE:
- Professional yet friendly
- Clear and concise in explanations
- Proactive in asking for clarification
- Transparent about your process
- Collaborative, not autonomous

CRITICAL FORMATTING RULES:
- You MAY use ## headers (two hashes) for major sections
- You MAY use ### headers (three hashes) for subsections
- NEVER use code blocks (triple backticks)
- NEVER use markdown lists with dashes (-) - use plain text with line breaks instead
- You MAY use **bold** for emphasis and *italic* for subtle emphasis
- Use plain text with line breaks for lists
- Use numbered format "1. " for numbered lists
- Format responses with simple line breaks and plain text structure
- When presenting compatibility scores, use clear format: "Compatibility Score: 0.72 (HIGH)"
- When showing data, use simple colon format: "Job Title: [title]" not markdown tables

**CRITICAL - File Downloads:**
- When you call generate_cv_pdf, generate_cv_docx, generate_cover_letter_pdf, or generate_cover_letter_docx tools, they return a success message
- You MUST include this EXACT success message in your response to the user
- DO NOT paraphrase or reword the tool's success message - copy it verbatim
- The success message contains the filename which is needed for the download system to work
- DO NOT create markdown links like [Download](sandbox:/filename.pdf) or [Download](file://filename.pdf)
- DO NOT create any file://, sandbox:/, or other file links - they don't work
- The user interface will AUTOMATICALLY display download buttons when it detects the tool's success messages
- If you generate BOTH formats (PDF and Word), include BOTH success messages and the UI will show TWO download buttons
- You can add additional text before or after the tool messages, but each tool's exact message MUST be included
- For Word documents, mention that formatting is minimal so users can customize it easily

Examples:
- Single format: Tool returns "CV PDF generated successfully! The file 'john_doe_cv.pdf' is ready for download."
  Your response: "Great! CV PDF generated successfully! The file 'john_doe_cv.pdf' is ready for download."
- Both formats: 
  Tool 1 returns: "CV PDF generated successfully! The file 'john_doe_cv.pdf' is ready for download."
  Tool 2 returns: "CV Word document generated successfully! The file 'john_doe_cv.docx' is ready for download."
  Your response MUST include BOTH messages: "Perfect! I've generated both formats for you. CV PDF generated successfully! The file 'john_doe_cv.pdf' is ready for download. CV Word document generated successfully! The file 'john_doe_cv.docx' is ready for download. You can customize the Word document formatting in Microsoft Word or Google Docs."

Remember: You are an assistant helping the user create their best application materials. 
The user is the expert on their own experience - you're the expert on presentation."""

# Create the Writer Agent
# This agent uses LangChain's create_agent which provides:
# - Automatic tool selection based on user input
# - Conversation memory across turns
# - Structured reasoning about when to use each tool
agent = create_agent(
    model=llm,
    tools=[
        analyze_cv_job_alignment,
        generate_tailored_cv_html,
        generate_cover_letter_content,
        generate_cv_pdf,
        generate_cover_letter_pdf,
        generate_cv_docx,
        generate_cover_letter_docx,
        # Enhanced tailoring tools
        generate_job_summary,
        calculate_semantic_similarity,
        calculate_bm25_score,
        calculate_compatibility_score,
        # v2 multi-dimensional scoring tools
        match_skill_pairs,
        assess_transferability_llm,
        calculate_compatibility_score_v2,
        build_gap_analysis,
        decide_tailoring_strategy,
        select_prioritize_content,
        rewrite_enhance_content
    ],
    system_prompt=system_prompt
)

# ============================================
# Interactive Runner (for testing)
# ============================================

def run_interactive_writer(cv_json: str = None, job_json: str = None, company_json: str = None):
    """
    Run writer agent interactively in terminal for testing.
    
    This function provides a CLI interface to test the writer agent with
    mock or real data from cv_agent and job_agent outputs.
    
    Args:
        cv_json: Optional ResumeInfo JSON string (if None, prompts user)
        job_json: Optional JobRequirements JSON string (if None, prompts user)
        company_json: Optional CompanyInfo JSON string (can be None)
        
    Returns:
        Dict with conversation history and results
        
    Example:
        >>> # With data
        >>> run_interactive_writer(cv_data, job_data)
        
        >>> # Interactive input
        >>> run_interactive_writer()
    """
    print("=" * 70)
    print("WRITER AGENT - Interactive Testing Mode")
    print("=" * 70)
    print("\nThis agent helps you tailor CVs and cover letters to job opportunities.")
    print("It works with structured JSON data from cv_agent and job_agent.\n")
    
    # Get data if not provided
    if not cv_json:
        print("Enter path to CV JSON file (or 'mock' for test data):")
        cv_input = input("> ").strip()
        if cv_input.lower() == 'mock':
            cv_json = json.dumps({
                "name": "Test Candidate",
                "email": "test@example.com",
                "phone": "+1234567890",
                "skills": ["Python", "JavaScript", "React"],
                "education": ["B.Sc. Computer Science"],
                "experience": ["Software Developer at Tech Corp"]
            })
        else:
            with open(cv_input, 'r') as f:
                cv_json = f.read()
    
    if not job_json:
        print("\nEnter path to Job JSON file (or 'mock' for test data):")
        job_input = input("> ").strip()
        if job_input.lower() == 'mock':
            job_json = json.dumps({
                "job_title": "Senior Software Engineer",
                "job_level": "Senior",
                "required_skills": ["Python", "React", "AWS"],
                "responsibilities": ["Build scalable applications"],
                "qualifications": ["Bachelor's degree in CS"]
            })
        else:
            with open(job_input, 'r') as f:
                job_json = f.read()
    
    # Initialize conversation
    initial_message = f"""I have CV and job data ready for tailoring.

CV Data:
{cv_json[:200]}...

Job Data:
{job_json[:200]}...

Please analyze the alignment and help me create a tailored CV."""
    
    messages = [{"role": "user", "content": initial_message}]
    
    # Conversation loop
    max_turns = 20
    turn = 0
    
    print("\n" + "=" * 70)
    print("Starting Conversation (type 'quit' to exit)")
    print("=" * 70)
    
    while turn < max_turns:
        turn += 1
        print(f"\n{'='*70}")
        print(f"Turn {turn}")
        print(f"{'='*70}\n")
        
        # Agent's turn
        print("🤖 Agent is processing...\n")
        try:
            result = agent.invoke({"messages": messages})
            agent_response = result["messages"][-1].content
            
            # Add to history
            messages.append({"role": "assistant", "content": agent_response})
            
            # Display response
            print(f"Agent:\n{agent_response}\n")
            
        except Exception as e:
            print(f"❌ Error: {e}")
            import traceback
            traceback.print_exc()
            break
        
        # User's turn
        print(f"{'-'*70}")
        user_input = input("You: ").strip()
        
        # Check for exit
        if user_input.lower() in ['quit', 'exit', 'q']:
            print("\n👋 Exiting writer agent.")
            break
        
        # Add user response
        messages.append({"role": "user", "content": user_input})
    
    if turn >= max_turns:
        print(f"\n⚠️  Reached maximum turns ({max_turns}).")
    
    return {
        "messages": messages,
        "turns": turn
    }

# ============================================
# CLI Entry Point
# ============================================

if __name__ == "__main__":
    import sys
    
    print("\n" + "="*70)
    print("WRITER AGENT - Command Line Interface")
    print("="*70 + "\n")
    
    # Check for command line arguments
    if len(sys.argv) > 2:
        cv_path = sys.argv[1]
        job_path = sys.argv[2]
        company_path = sys.argv[3] if len(sys.argv) > 3 else None
        
        print(f"Loading CV data from: {cv_path}")
        print(f"Loading Job data from: {job_path}")
        if company_path:
            print(f"Loading Company data from: {company_path}")
        
        try:
            with open(cv_path, 'r') as f:
                cv_json = f.read()
            with open(job_path, 'r') as f:
                job_json = f.read()
            company_json = None
            if company_path:
                with open(company_path, 'r') as f:
                    company_json = f.read()
            
            run_interactive_writer(cv_json, job_json, company_json)
            
        except FileNotFoundError as e:
            print(f"\n❌ Error: File not found - {e}")
            sys.exit(1)
        except Exception as e:
            print(f"\n❌ Error: {e}")
            import traceback
            traceback.print_exc()
            sys.exit(1)
    else:
        print("Usage:")
        print("  python writer_agent.py <cv_json_path> <job_json_path> [company_json_path]")
        print("\nOr run without arguments for interactive mode:")
        print("  python writer_agent.py\n")
        
        proceed = input("Run in interactive mode? (y/n): ").strip().lower()
        if proceed == 'y':
            try:
                run_interactive_writer()
            except KeyboardInterrupt:
                print("\n\n👋 Interrupted by user.")
            except Exception as e:
                print(f"\n❌ Error: {e}")
                import traceback
                traceback.print_exc()
        else:
            print("Exiting.")
